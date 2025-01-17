# quality_monitoring/src/quality_monitoring/app.py

import streamlit as st
from pathlib import Path
import json
import time
from typing import List
import yaml
from quality_monitoring.crew import QualityMonitoring
from dotenv import load_dotenv
from openai import OpenAI
import markdown
from weasyprint import HTML
from docx import Document
from io import BytesIO
import tempfile

load_dotenv()

client = OpenAI()

class QualityMonitorApp:
    def __init__(self):
        self.setup_directories()
        self.initialize_session_state()
        self.setup_page_config()
        
    def setup_directories(self):
        """Setup necessary directories"""
        self.config_path = Path(__file__).parent / "config"
        self.output_path = Path(__file__).parent.parent.parent / "output"
        self.feedback_path = Path(__file__).parent.parent.parent / "feedback"
        
        self.output_path.mkdir(exist_ok=True)
        self.feedback_path.mkdir(exist_ok=True)
        
        self.languages = {
            "🇺🇸 English": "en",
            "🇧🇷 Português": "pt",
            "🇪🇸 Español": "es"
        }

    def initialize_session_state(self):
        """Initialize session state variables"""
        if 'analysis_result' not in st.session_state:
            st.session_state.analysis_result = None
        if 'current_transcript' not in st.session_state:
            st.session_state.current_transcript = ""
        if 'current_guidelines' not in st.session_state:
            st.session_state.current_guidelines = self.load_default_guidelines()

    def setup_page_config(self):
        """Setup page configuration"""
        st.set_page_config(
            page_title="Customer Service Quality Monitor",
            page_icon="🎯",
            layout="wide",
            initial_sidebar_state="expanded"
        )

    def set_custom_styles(self):
        """Apply custom CSS styles"""
        st.markdown("""
        <style>
        .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        .main-header {
            background: linear-gradient(to right, #3498db, #2ecc71);
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            color: white;
            text-align: center;
        }
        .stTextArea textarea {
            background-color: #2d3436;
            color: #dfe6e9;
            border: 1px solid #4a5568;
        }
        .stButton button {
            width: 100%;
            padding: 0.5rem;
            background: linear-gradient(to right, #3498db, #2ecc71);
            color: white;
            border: none;
            border-radius: 5px;
        }
        .results-card {
            background-color: #2d3436;
            padding: 1.5rem;
            border-radius: 10px;
            margin-top: 1rem;
        }
        .stMarkdown {
            color: #dfe6e9;
        }
        .feedback-card {
            background-color: #2d3436;
            padding: 1.5rem;
            border-radius: 10px;
            margin-top: 1rem;
        }
        </style>
        """, unsafe_allow_html=True)

    def load_default_guidelines(self) -> List[str]:
        return [
            "Greet the customer warmly",
            "Listen actively and show empathy",
            "Provide clear and helpful solutions",
            "Confirm customer satisfaction",
            "Offer additional assistance"
        ]

    def save_feedback(self, transcript: str, guidelines: List[str], rating: int, feedback_text: str, language: str, analysis_result: str):
        feedback = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'transcript': transcript,
            'guidelines': guidelines,
            'rating': rating,
            'feedback': feedback_text,
            'language': language,
            'analysis_result': analysis_result,
            'metadata': {
                'version': '1.0',
                'feedback_type': 'customer_service_quality',
                'timestamp_utc': time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
            }
        }
        
        # Ensure the feedback directory exists
        self.feedback_path.mkdir(exist_ok=True)
        
        # Create a timestamped filename with additional info
        timestamp = time.strftime('%Y%m%d_%H%M%S')
        feedback_file = self.feedback_path / f'feedback_{timestamp}_rating{rating}.json'
        
        # Save feedback with pretty printing
        with open(feedback_file, 'w', encoding='utf-8') as f:
            json.dump(feedback, f, ensure_ascii=False, indent=2)

    def translate_final_output(self, output: str, language: str) -> str:
        """Translate the final output to the selected language"""
        if language == 'en':
            return output
        
        lang_name = "portuguese" if language == 'pt' else "spanish"
        
        formated_output = "Translate to " + lang_name + ":\n" + str(output)
        try:

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert translator. The user will only give you the text and the language they want it translated to. Your answer should be only the exact same text translated to the requested language following the same format. Answer formatted as markdown without '```'"},
                    {"role": "user", "content": formated_output}
                ],
                temperature=0.2,
                stream=False,
            )
            return response.choices[0].message.content
        
        except Exception as e:
            return f"Error translating output: {str(e)}"

    def run_analysis(self, transcript: str, guidelines: List[str], language: str) -> str:
        inputs = {
            'transcript': transcript,
            'guidelines': guidelines,
            'language': language
        }
        
        crew = QualityMonitoring(inputs)
        result = crew.crew().kickoff(inputs=inputs)

        if language != 'en':
            translation = self.translate_final_output(result, language)

        if not translation.startswith("Error"):
            result = translation
            
        return result

    def render_sidebar(self):
        """Render sidebar content"""
        with st.sidebar:
            st.markdown("## Settings")
            
            language = st.selectbox(
                "Analysis Language",
                options=list(self.languages.keys()),
                index=0,
                help="Select the language for the analysis output"
            )

            st.markdown("## About")
            st.markdown("""
            This tool helps evaluate customer service quality using AI.
            
            **Features:**
            - Multi-language support
            - Real-time analysis
            - Detailed feedback
            - Quality metrics
            """)

            st.markdown("## Help")
            with st.expander("How to use"):
                st.markdown("""
                1. Paste your transcript
                2. Review guidelines or provide your own separated by new lines
                3. Select language
                4. Click Analyze
                5. Rate the results
                """)

        return self.languages[language]

    def markdown_to_pdf(self, markdown_text: str) -> bytes:
        """Convert markdown text to PDF"""
        # Convert markdown to HTML
        html = markdown.markdown(markdown_text)
        
        # Create a temporary file for the PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            # Convert HTML to PDF using WeasyPrint
            HTML(string=f"""
                <html>
                    <head>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 2cm; }}
                            h1, h2, h3 {{ color: #2c3e50; }}
                        </style>
                    </head>
                    <body>{html}</body>
                </html>
            """).write_pdf(tmp.name)
            
            # Read the generated PDF
            with open(tmp.name, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
        
        return pdf_content
    
    def markdown_to_docx(self, markdown_text: str) -> bytes:
        """Convert markdown text to DOCX"""
        doc = Document()
        
        # Simple markdown parsing (you might want to use a more robust parser)
        lines = markdown_text.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], 0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
            elif line:
                doc.add_paragraph(line)
        
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()

    def render_download_buttons(self):
        """Render download buttons for PDF and DOCX"""
        col1, col2 = st.columns(2)
        
        with col1:
            pdf_content = self.markdown_to_pdf(st.session_state.analysis_result)
            st.download_button(
                label="📥 Download as PDF",
                data=pdf_content,
                file_name="analysis_report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        
        with col2:
            docx_content = self.markdown_to_docx(st.session_state.analysis_result)
            st.download_button(
                label="📥 Download as DOCX",
                data=docx_content,
                file_name="analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    def render_main_content(self, language_code: str):
        """Render main content area"""
        st.markdown(
            '<div class="main-header">'
            '<h1>🎯 Customer Service Quality Monitor</h1>'
            '<p>Evaluate and improve your customer service interactions using AI</p>'
            '</div>',
            unsafe_allow_html=True
        )

        # Main content area
        col1, col2 = st.columns([3, 2])
        
        with col1:
            st.markdown("### Customer Service Transcript")
            transcript = st.text_area(
                label="Service Transcript",
                value=st.session_state.current_transcript,
                height=300,
                placeholder="Paste your customer service transcript here...",
                help="Paste the conversation between customer service representative and customer",
                key="transcript_input",
                label_visibility="collapsed"
            )

        with col2:
            st.markdown("### Service Guidelines")
            guidelines_text = st.text_area(
                label="Service Guidelines",
                value="\n".join(st.session_state.current_guidelines),
                height=300,
                help="These guidelines will be used to evaluate the service quality",
                key="guidelines_input",
                label_visibility="collapsed"
            )
            guidelines = [g.strip() for g in guidelines_text.split("\n") if g.strip()]

        # Analysis button
        if st.button("🔍 Analyze Interaction", use_container_width=True, type="primary"):
            if not transcript:
                st.error("Please provide a transcript to analyze")
                return
                
            with st.spinner("🔄 Analyzing transcript..."):
                try:
                    result = self.run_analysis(transcript, guidelines, language_code)
                    st.session_state.analysis_result = result
                    st.session_state.current_transcript = transcript
                    st.session_state.current_guidelines = guidelines
                    
                except Exception as e:
                    st.error(f"An error occurred during analysis: {str(e)}")
                    return

        # Display results if available
        if st.session_state.analysis_result:
            st.markdown('<div class="results-card">', unsafe_allow_html=True)
            st.success("✅ Analysis complete!")
            st.markdown("### Analysis Results")
            st.markdown(st.session_state.analysis_result)
            
            # Add download buttons
            st.markdown("### Download Report")
            self.render_download_buttons()
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Feedback section
            st.markdown('<div class="feedback-card">', unsafe_allow_html=True)
            st.markdown("### Rate This Analysis")
            col1, col2 = st.columns(2)
            
            with col1:
                rating = st.slider(
                    label="Analysis Rating",
                    min_value=1,
                    max_value=5,
                    value=3,
                    help="1 = Not helpful, 5 = Very helpful",
                    key="rating_slider"
                )
            
            with col2:
                feedback_text = st.text_area(
                    label="Feedback Comments",
                    placeholder="Share your thoughts about the analysis...",
                    key="feedback_text"
                )
            
            if st.button("📮 Submit Feedback", key="submit_feedback"):
                self.save_feedback(
                    st.session_state.current_transcript,
                    st.session_state.current_guidelines,
                    rating,
                    feedback_text,
                    language_code,
                    st.session_state.analysis_result
                )
                st.success("Thank you for your feedback! 🙏")
            
            st.markdown('</div>', unsafe_allow_html=True)

    def run_app(self):
        """Main Streamlit application"""
        self.set_custom_styles()
        language_code = self.render_sidebar()
        self.render_main_content(language_code)

def main():
    app = QualityMonitorApp()
    app.run_app()

if __name__ == "__main__":
    main()